import io
from pathlib import Path
from typing import Optional

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document


def read_text_from_file(file_path: Path) -> str:
    """
    Read and extract text from a resume file. Supports PDF, DOCX, and TXT.
    """
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return pdf_extract_text(str(file_path)) or ""
    if suffix in {".docx"}:
        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs)
    if suffix in {".txt"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")


def read_text_from_bytes(name: str, data: bytes) -> str:
    """
    Extract text when given an uploaded file's name and bytes content.
    """
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        with io.BytesIO(data) as fp:
            return pdf_extract_text(fp) or ""
    if suffix in {".docx"}:
        with io.BytesIO(data) as fp:
            doc = Document(fp)
            return "\n".join(p.text for p in doc.paragraphs)
    if suffix in {".txt"}:
        return data.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix}")



